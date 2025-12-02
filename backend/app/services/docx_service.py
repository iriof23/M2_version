"""
DOCX Service - Word Document Generation Engine

Renders "Editable" Word documents from templates using docxtpl.
Uses plain text versions of rich text fields for maximum compatibility.
"""
import io
import logging
import re
from pathlib import Path
from typing import Optional, Any, Union
from datetime import datetime

from docxtpl import DocxTemplate, RichText, Listing
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.report_data_service import ReportContext, ReportDataService
from app.services.rich_text_service import RichTextService

logger = logging.getLogger(__name__)

# Template directory path
TEMPLATES_DIR = Path(__file__).parent.parent / "templates" / "docx"


class DOCXService:
    """
    DOCX Generation Service using docxtpl.
    
    Renders Word document templates with report data.
    Uses plain text versions of rich text fields for stability.
    """
    
    # Severity color mapping for Word (RGB tuples)
    SEVERITY_COLORS = {
        'CRITICAL': RGBColor(220, 38, 38),    # Red
        'HIGH': RGBColor(234, 88, 12),        # Orange
        'MEDIUM': RGBColor(202, 138, 4),      # Yellow/Amber
        'LOW': RGBColor(37, 99, 235),         # Blue
        'INFO': RGBColor(107, 114, 128),      # Gray
        'INFORMATIONAL': RGBColor(107, 114, 128),
    }

    def __init__(self):
        """Initialize the DOCX service."""
        # Ensure templates directory exists
        TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

    async def generate_docx(
        self,
        report_id: str,
        output_path: Optional[str] = None,
        template_name: str = "master.docx",
    ) -> bytes:
        """
        Generate a DOCX report from the given report ID.
        
        Args:
            report_id: UUID of the report to generate
            output_path: Optional path to save the DOCX file
            template_name: Name of the template file to use
            
        Returns:
            DOCX content as bytes
            
        Raises:
            ValueError: If report not found or template missing
            RuntimeError: If DOCX generation fails
        """
        logger.info(f"Generating DOCX for report: {report_id}")
        
        # Build context from report data
        context = await ReportDataService.build_context(report_id)
        
        # Convert to dict for docxtpl
        context_dict = ReportDataService.context_to_dict(context)
        
        # Generate DOCX from context
        docx_bytes = await self.generate_docx_from_context(
            context_dict,
            template_name=template_name
        )
        
        # Save to file if path provided
        if output_path:
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'wb') as f:
                f.write(docx_bytes)
            logger.info(f"DOCX saved to: {output_path}")
        
        return docx_bytes

    async def generate_docx_from_context(
        self,
        context: dict,
        template_name: str = "master.docx",
    ) -> bytes:
        """
        Generate DOCX from a pre-built context dictionary.
        
        Args:
            context: Template context dictionary
            template_name: Name of the template file
            
        Returns:
            DOCX content as bytes
        """
        template_path = TEMPLATES_DIR / template_name
        
        # Check if template exists, otherwise use fallback
        if not template_path.exists():
            logger.warning(f"Template not found: {template_path}, using fallback")
            return self._generate_fallback_docx(context)
        
        try:
            # Load template
            doc = DocxTemplate(str(template_path))
            
            # Process context for Word compatibility
            processed_context = self._process_context_for_docx(context, doc)
            
            # Render template
            doc.render(processed_context)
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            docx_bytes = buffer.getvalue()
            logger.info(f"Generated DOCX: {len(docx_bytes)} bytes")
            
            return docx_bytes
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            raise RuntimeError(f"DOCX generation failed: {e}")

    def _process_context_for_docx(self, context: dict, doc: DocxTemplate) -> dict:
        """
        Process context dictionary for Word compatibility.
        
        - Converts HTML fields to plain text
        - Creates RichText objects for styled content
        - Handles special formatting
        
        Args:
            context: Raw context dictionary
            doc: DocxTemplate instance for creating subdocs
            
        Returns:
            Processed context dictionary
        """
        processed = context.copy()
        
        # Process findings - use plain text versions
        if 'findings' in processed:
            processed_findings = []
            for finding in processed['findings']:
                processed_finding = finding.copy()
                
                # Use plain text for Word stability
                processed_finding['description'] = finding.get('description_plain', '')
                processed_finding['remediation'] = finding.get('remediation_plain', '')
                
                # Create styled severity text
                severity = finding.get('severity', 'MEDIUM').upper()
                processed_finding['severity_styled'] = self._create_severity_richtext(severity)
                
                # Format references as bullet list
                refs = finding.get('references', [])
                if refs:
                    processed_finding['references_text'] = '\n'.join(f'• {ref}' for ref in refs)
                else:
                    processed_finding['references_text'] = 'N/A'
                
                processed_findings.append(processed_finding)
            
            processed['findings'] = processed_findings
        
        # Process executive summary - use plain text
        if 'executive_summary_html' in processed:
            processed['executive_summary'] = processed.get('executive_summary_plain', '')
        
        # Process project description
        if 'project' in processed:
            project = processed['project'].copy()
            project['description'] = project.get('description_plain', '')
            
            # Format scope as bullet list
            scope = project.get('scope', [])
            if scope:
                project['scope_text'] = '\n'.join(f'• {item}' for item in scope)
            else:
                project['scope_text'] = 'No scope items defined'
            
            # Format compliance frameworks
            frameworks = project.get('compliance_frameworks', [])
            if frameworks:
                project['compliance_text'] = ', '.join(frameworks)
            else:
                project['compliance_text'] = 'N/A'
            
            processed['project'] = project
        
        # Add generation metadata
        processed['generation_date'] = datetime.now().strftime("%B %d, %Y")
        processed['generation_time'] = datetime.now().strftime("%H:%M:%S")
        
        return processed

    def _create_severity_richtext(self, severity: str) -> RichText:
        """
        Create a styled RichText object for severity badge.
        
        Args:
            severity: Severity level string
            
        Returns:
            RichText object with colored severity text
        """
        rt = RichText()
        color = self.SEVERITY_COLORS.get(severity.upper(), RGBColor(107, 114, 128))
        rt.add(severity, bold=True, color=color, size=Pt(11))
        return rt

    def _generate_fallback_docx(self, context: dict) -> bytes:
        """
        Generate a basic DOCX document when no template is available.
        
        Creates a simple but professional report structure.
        
        Args:
            context: Report context dictionary
            
        Returns:
            DOCX content as bytes
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = context.get('report_title', 'Security Assessment Report')
        doc.core_properties.author = context.get('generated_by', 'Security Team')
        
        # Title
        title = doc.add_heading(context.get('report_title', 'Security Assessment Report'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Client: {context.get('client', {}).get('name', 'N/A')}\n").bold = True
        meta_para.add_run(f"Project: {context.get('project', {}).get('name', 'N/A')}\n")
        meta_para.add_run(f"Date: {context.get('generated_at', 'N/A')}\n")
        meta_para.add_run(f"Classification: {context.get('classification', 'CONFIDENTIAL')}")
        
        doc.add_page_break()
        
        # Executive Summary
        doc.add_heading('1. Executive Summary', 1)
        
        stats = context.get('stats', {})
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"Total Findings: ").bold = True
        summary_para.add_run(f"{stats.get('total_findings', 0)}\n")
        summary_para.add_run(f"Critical: ").bold = True
        summary_para.add_run(f"{stats.get('critical_count', 0)}  ")
        summary_para.add_run(f"High: ").bold = True
        summary_para.add_run(f"{stats.get('high_count', 0)}  ")
        summary_para.add_run(f"Medium: ").bold = True
        summary_para.add_run(f"{stats.get('medium_count', 0)}  ")
        summary_para.add_run(f"Low: ").bold = True
        summary_para.add_run(f"{stats.get('low_count', 0)}\n")
        summary_para.add_run(f"Risk Level: ").bold = True
        summary_para.add_run(f"{stats.get('risk_level', 'N/A')}")
        
        # Executive summary content
        if context.get('executive_summary_plain'):
            doc.add_paragraph(context['executive_summary_plain'])
        
        # Scope & Methodology
        doc.add_heading('2. Scope & Methodology', 1)
        
        project = context.get('project', {})
        scope_para = doc.add_paragraph()
        scope_para.add_run(f"Assessment Type: ").bold = True
        scope_para.add_run(f"{project.get('project_type', 'Penetration Test')}\n")
        scope_para.add_run(f"Methodology: ").bold = True
        scope_para.add_run(f"{project.get('methodology', 'OWASP Testing Guide')}\n")
        scope_para.add_run(f"Testing Period: ").bold = True
        scope_para.add_run(f"{project.get('start_date', 'N/A')} - {project.get('end_date', 'N/A')}")
        
        # Scope items
        scope_items = project.get('scope', [])
        if scope_items:
            doc.add_heading('Scope Items:', 2)
            for item in scope_items:
                doc.add_paragraph(item, style='List Bullet')
        
        doc.add_page_break()
        
        # Findings Summary Table
        doc.add_heading('3. Findings Summary', 1)
        
        findings = context.get('findings', [])
        if findings:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'ID'
            header_cells[1].text = 'Title'
            header_cells[2].text = 'Severity'
            header_cells[3].text = 'CVSS'
            
            for cell in header_cells:
                cell.paragraphs[0].runs[0].bold = True
            
            # Data rows
            for finding in findings:
                row = table.add_row().cells
                row[0].text = finding.get('reference_id', 'N/A')
                row[1].text = finding.get('title', 'Untitled')
                row[2].text = finding.get('severity', 'N/A')
                row[3].text = str(finding.get('cvss_score', 'N/A'))
        else:
            doc.add_paragraph('No findings reported.')
        
        doc.add_page_break()
        
        # Detailed Findings
        doc.add_heading('4. Detailed Findings', 1)
        
        for idx, finding in enumerate(findings, 1):
            severity = finding.get('severity', 'MEDIUM')
            title = finding.get('title', 'Untitled Finding')
            
            # Finding header
            finding_heading = doc.add_heading(f"{idx}. [{severity}] {title}", 2)
            
            # Finding ID
            id_para = doc.add_paragraph()
            id_para.add_run(f"Reference ID: ").bold = True
            id_para.add_run(finding.get('reference_id', 'N/A'))
            
            # Description
            doc.add_heading('Description', 3)
            desc = finding.get('description_plain') or finding.get('description', 'No description provided.')
            doc.add_paragraph(desc)
            
            # Remediation
            remediation = finding.get('remediation_plain') or finding.get('remediation', '')
            if remediation:
                doc.add_heading('Remediation', 3)
                doc.add_paragraph(remediation)
            
            # Metadata
            meta_para = doc.add_paragraph()
            if finding.get('cvss_score'):
                meta_para.add_run(f"CVSS Score: ").bold = True
                meta_para.add_run(f"{finding['cvss_score']}  ")
            if finding.get('cve_id'):
                meta_para.add_run(f"CVE: ").bold = True
                meta_para.add_run(f"{finding['cve_id']}  ")
            if finding.get('affected_systems'):
                meta_para.add_run(f"Affected Systems: ").bold = True
                meta_para.add_run(finding['affected_systems'])
            
            doc.add_paragraph()  # Spacer
        
        # Conclusion
        doc.add_heading('5. Conclusion', 1)
        conclusion = doc.add_paragraph()
        conclusion.add_run(
            f"This assessment identified {stats.get('total_findings', 0)} vulnerabilities. "
        )
        if stats.get('critical_count', 0) > 0:
            conclusion.add_run(
                f"The {stats['critical_count']} critical findings require immediate attention. "
            )
        conclusion.add_run(
            "We recommend prioritizing remediation based on severity and scheduling "
            "a follow-up assessment to verify implemented controls."
        )
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.add_run("─" * 50 + "\n")
        footer_para.add_run(
            f"This report is confidential and intended solely for {context.get('client', {}).get('name', 'the client')}."
        )
        
        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()

    def create_findings_table(self, findings: list, doc: DocxTemplate) -> Any:
        """
        Create a formatted findings table for insertion into template.
        
        Can be used with docxtpl's {{findings_table}} placeholder.
        
        Args:
            findings: List of finding dictionaries
            doc: DocxTemplate instance
            
        Returns:
            Subdocument or table object for template insertion
        """
        # This would create a subdoc with a table
        # For MVP, we use the Listing approach in the template directly
        pass


# Singleton instance
docx_service = DOCXService()

